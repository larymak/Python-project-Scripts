""" This module uses 'python-docx' package to analyze a Word Document ('.docx')
    This module can be imported in other programs but the preffered way is to run it 
    as a Script from the command line like this:
    
    $ python extract.py <name-of-word-document-file>
    
    The basic function of the program is to analyze the word document and 
    collect bold, italic and underlined words from it and then 
    after analyzing write these collected words at the end of the word document.
    So the program first reads the word document, collects bold, italic 
    and underlined words from it and then writes the collected words at
    the very end of the same word document 
    Copyright 2023 Kashaan Mahmood
    License: MIT License
             https://opensource.org/license/mit/
    """


from docx import Document
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# global variables
total_words = 0
wordsList = ""


# calculate total words in docx
def get_total_words(docxFile):
    document = Document(docxFile)
    total = 0
    for p in document.paragraphs:
        for run in p.runs:
            total += len(run.text.split())
    return total


unwanted_characters = [
    '"',
    "'",
    "’",
    "“",
    ":",
    "\n",
    "-",
    "— — ",
    "—",
    ".",
    ",",
    ";",
    "!",
    "?",
]


def remove_unwanted(words):
    """remove unwanted characters from analyzed output"""
    for i in unwanted_characters:
        if i in words:
            words = words.replace(i, "")
    return words


def analyze(docxFile):
    """analyze the docx file and collect bold, italicized, and underlined words from it
    and return a `collect` object these selected words
    """

    document = Document(docxFile)

    collect = [
        {"b": []},
        {"i": []},
        {"u": []},
        {"bi": []},
        {"bu": []},
        {"iu": []},
        {"biu": []},
    ]

    for p in document.paragraphs:
        for run in p.runs:
            if run.bold and run.italic and run.underline:
                filtered_text = remove_unwanted(run.text)
                collect[6]["biu"].append(filtered_text)

            elif run.bold and run.italic:
                filtered_text = remove_unwanted(run.text)
                collect[3]["bi"].append(filtered_text)

            elif run.bold and run.underline:
                filtered_text = remove_unwanted(run.text)
                collect[4]["bu"].append(filtered_text)

            elif run.italic and run.underline:
                filtered_text = remove_unwanted(run.text)
                collect[5]["iu"].append(filtered_text)

            elif run.bold:
                filtered_text = remove_unwanted(run.text)
                collect[0]["b"].append(filtered_text)

            elif run.italic:
                filtered_text = remove_unwanted(run.text)
                collect[1]["i"].append(filtered_text)

            elif run.underline:
                filtered_text = remove_unwanted(run.text)
                collect[2]["u"].append(filtered_text)

    return collect


def write_data(docxFile, data):
    """gets the `collect` variable as 'data' argument from analyze()
    and reads and appends the 'data' to end of docx file
    """

    global wordsList

    document = Document(docxFile)

    def save_document():
        document.save(docxFile)
        return "saved"

    def add_words(key):
        global wordsList
        categories = {
            "b": "\nBold Words:-",
            "i": "\n\nItalicized Words:-",
            "u": "\n\nUnderlined Words:-",
            "bi": "\n\nBold & Italicized Words:-",
            "bu": "\n\nBold & Underlined Words:-",
            "biu": "\n\nBold & Italicized & Underlined Words:-",
            "iu": "\n\nItalicized & Underlined Words:-",
        }
        for word in words[key]:
            category = categories[key]
            if len(wordsList) == 0 or category not in wordsList:
                wordsList = wordsList + f"{category}\n{word}"
            else:
                wordsList = wordsList + f", {word}"

    title_p = document.add_paragraph(f"\n========== Extracted Words ==========\n")

    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    write_p = document.add_paragraph()

    for words in data:
        if words.__contains__("b") and words["b"]:
            add_words("b")

        elif words.__contains__("u") and words["u"]:
            add_words("u")

        elif words.__contains__("bi") and words["bi"]:
            add_words("bi")

        elif words.__contains__("bu") and words["bu"]:
            add_words("bu")

        elif words.__contains__("iu") and words["iu"]:
            add_words("iu")

        elif words.__contains__("biu") and words["biu"]:
            add_words("biu")

    write_p.add_run(f"{wordsList}")

    ending_p = document.add_paragraph("\n===================\n")
    ending_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    save_document()
    return


# function calls inside main()


def main():
    global total_words

    data = analyze(docx)
    write_data(docx, data)


if __name__ == "__main__":
    from sys import argv
    import time

    # get docx file
    docx = argv[1]

    print(f"Started at {time.strftime('%X')}...")
    # calling main()
    main()
    print(f"Finished at {time.strftime('%X')}...")
